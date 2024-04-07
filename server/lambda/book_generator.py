import os
import re
import time
import base64
import requests
import json
from docx import Document
from docx.shared import RGBColor, Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Define your API keys here
ANTHROPIC_API_KEY = "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"
# ANTHROPIC_API_KEY = "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"
stability_api_key = "sk-0tRzR3toAMfTVDs0OH9FnSUj0wtR5DiV5oUjt1VrrcuW30UT"


def remove_first_line(test_string):
    return re.sub(r'^.*\n', '', test_string, count=1) if test_string.startswith("Here") and test_string.split("\n")[0].strip().endswith(":") else test_string


def word_count(s):
    return len(re.findall(r'\w+', s))


def generate_text(prompt, model="claude-3-haiku-20240307", max_tokens=3000, temperature=0.7, retries=5, max_wait=60):
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    data = {
        "model": model,
        "max_tokens": max_tokens,
        "temperature": temperature,
        "system": "You are a world-class author. Write the requested content with great skill and attention to detail.",
        "messages": [{"role": "user", "content": prompt}],
    }
    for attempt in range(retries):
        response = requests.post(
            "https://api.anthropic.com/v1/messages", headers=headers, json=data)
        if response.status_code == 200:
            response_json = response.json()
            try:
                response_text = response_json['content'][0]['text'].strip()
                if not response_text:
                    raise ValueError("Empty response from the model.")
                return response_text
            except (KeyError, IndexError):
                raise ValueError(
                    "Unexpected response structure from the model.")
        elif response.status_code == 429:
            wait = min(2 ** attempt, max_wait)
            print(f"Rate limit exceeded. Retrying in {wait} seconds...")
            time.sleep(wait)
        else:
            raise ValueError(
                f"Failed to fetch data from API. Status Code: {response.status_code}, Attempt: {attempt+1}")
    raise Exception("Max retries exceeded.")


def generate_cover_prompt(plot):
    prompt = f"Plot: {plot}\n\n--\n\nDescribe the cover we should create, based on the plot. This should be visually rich and detailed, ideally two sentences long."
    generated_prompt = generate_text(prompt)
    if not generated_prompt.strip():
        raise ValueError("Generated prompt is empty.")
    if len(generated_prompt) > 1000:
        raise ValueError("Generated prompt is too long.")
    if "example_disallowed_content" in generated_prompt.lower():
        raise ValueError("Generated prompt contains disallowed content.")
    return generated_prompt


def generate_title(plot):
    prompt = f"Here is the plot for the book: {plot}\n\n--\n\nRespond with a great title for this book. Only respond with the title, nothing else is allowed."
    return remove_first_line(generate_text(prompt))


def create_cover_image(plot_description):
    engine_id = "stable-diffusion-xl-beta-v2-2-2"
    api_host = 'https://api.stability.ai'
    api_key = stability_api_key  # Corrected variable name
    response = requests.post(
        f"{api_host}/v1/generation/{engine_id}/text-to-image",
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Authorization": f"Bearer {api_key}"
        },
        json={
            "text_prompts": [{"text": plot_description}],
            "cfg_scale": 7,
            "clip_guidance_preset": "FAST_BLUE",
            "height": 768,
            "width": 512,
            "samples": 1,
            "steps": 30,
        },
    )
    # The rest of your function continues...

    if response.status_code != 200:
        raise Exception("Non-200 response: " + str(response.text))
    data = response.json()
    image_path = "./cover.png"  # Changed to a local file path
    for image in data["artifacts"]:
        with open(image_path, "wb") as f:
            f.write(base64.b64decode(image["base64"]))
    return image_path


def create_doc(title, author, chapters, chapter_titles, cover_image_path, file_stream):
    document = Document()

    try:
        document.add_picture(cover_image_path, width=Inches(6))
        document.add_page_break()
    except Exception as e:
        print(f"Failed to add cover image: {e}. Continuing without it.")

    document.add_heading(title, level=0)
    run = document.add_paragraph().add_run(f"Author: {author}\n\n")
    run.bold = True

    for i, chapter_title in enumerate(chapter_titles):
        chapter_content = chapters[i]

        document.add_heading(chapter_title, level=1)
        paragraphs = chapter_content.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                run = document.add_paragraph().add_run(paragraph.strip())
                if paragraph.startswith("Subtitle"):
                    run.bold = True
                    run.font.color.rgb = RGBColor(34, 139, 34)

    document.save(file_stream)


def generate_book(writing_style, book_description, chapter_titles, chapter_elaborations):
    all_chapters_content = []

    for i, chapter_title in enumerate(chapter_titles):
        print(f"Generating content for chapter '{chapter_title}'...")
        subtitles_content = []

        # Let's assume each chapter should have content for 8 subtitles
        for subtitle_index in range(1, 9):
            attempts, subtitle_generated, max_attempts = 0, False, 3

            while not subtitle_generated and attempts < max_attempts:
                subtitle_prompt = f"Subtitle {subtitle_index}: Provide a detailed analysis and insights for '{chapter_title}', as part of a book in the style of '{writing_style}'. Ensure this section contains at least 800 words."
                
                # Include the chapter elaboration in the prompt if it's not empty
                if chapter_elaborations[i]:
                    subtitle_prompt += f" Follow the instructions provided and include the following additional information after thinking it through intelligently. Add them in a way that is balanced and not excessive or abnormal: {chapter_elaborations[i]}"
                
                subtitle_content = generate_text(subtitle_prompt, max_tokens=3000)

                # Ensuring content meets the word count
                if word_count(subtitle_content) >= 500:
                    subtitle_generated = True
                    subtitles_content.append(subtitle_content)
                else:
                    attempts += 1
                    print(
                        f"Insufficient content for subtitle {subtitle_index} in chapter '{chapter_title}'. Attempting regeneration, attempt #{attempts}.")

        # Combine all subtitles into one chapter content
        chapter_content = "\n\n".join(subtitles_content)
        all_chapters_content.append(chapter_content)
        time.sleep(1)  # Pause between chapters to avoid rate limiting

    print("Book content generation completed.")
    return all_chapters_content