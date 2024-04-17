import os
import re
import time
import requests
import json
import uuid
import io
import pymongo

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bson.binary import Binary
from bson.objectid import ObjectId
from pymongo import MongoClient
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Environment variable for MongoDB URI
MONGO_URI = 'mongodb+srv://aba326ss:160160ssSS@cluster0.2h1ad.mongodb.net/ai-books?retryWrites=true&w=majority'  # Make sure to set this in your .env file

ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')

# Initialize MongoDB connection
client = MongoClient(MONGO_URI)
db = client.get_default_database()  # This will connect to the default database specified in the URI
collection = db['books'] 

def remove_first_line(test_string):
    return re.sub(r'^.*\n', '', test_string, count=1) if test_string.startswith("Here") and test_string.split("\n")[0].strip().endswith(":") else test_string

def word_count(s):
    return len(re.findall(r'\w+', s))

def generate_text(prompt, model="claude-3-haiku-20240307", max_tokens=3000, temperature=0.7, retries=5, max_wait=60):
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    data = {
        "model": model,
        "max_tokens": max_tokens,
        "temperature": temperature,
        "system": "You are a world-class author. Write the requested content with great skill and attention to detail.",
        "messages": [{"role": "user", "content": prompt}],
    }
    for attempt in range(retries):
        response = requests.post(
            "https://api.anthropic.com/v1/messages", headers=headers, json=data)
        if response.status_code == 200:
            response_json = response.json()
            try:
                response_text = response_json['content'][0]['text'].strip()
                if not response_text:
                    raise ValueError("Empty response from the model.")
                return response_text
            except (KeyError, IndexError):
                raise ValueError("Unexpected response structure from the model.")
        elif response.status_code == 429:
            wait = min(2 ** attempt, max_wait)
            print(f"Rate limit exceeded. Retrying in {wait} seconds...")
            time.sleep(wait)
        else:
            raise ValueError(f"Failed to fetch data from API. Status Code: {response.status_code}, Attempt: {attempt+1}")
    raise Exception("Max retries exceeded.")

def generate_title(plot):
    prompt = f"Here is the plot for the book: {plot}\n\n--\n\nRespond with a great title for this book. Only respond with the title, nothing else is allowed."
    return remove_first_line(generate_text(prompt))

def create_doc(title, author, chapters, chapter_titles, file_stream):
    document = Document()
    document.add_heading(title, level=0)
    run = document.add_paragraph().add_run(f"Author: {author}\n\n")
    run.bold = True

    for i, chapter_title in enumerate(chapter_titles):
        chapter_content = chapters[i]
        document.add_heading(chapter_title, level=1)
        paragraphs = chapter_content.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                run = document.add_paragraph().add_run(paragraph.strip())
                if paragraph.startswith("Subtitle"):
                    run.bold = True
                    run.font.color.rgb = RGBColor(34, 139, 34)

    document.save(file_stream)

def save_book(file_stream, title, unique_filename):
    # Convert file stream to binary
    binary_file = Binary(file_stream.getvalue())
    
    # Create a document to store in MongoDB
    book_document = {
        "filename": unique_filename,
        "title": title,
        "content": binary_file
    }
    
    # Insert the document into the collection
    result = collection.insert_one(book_document)
    
    # Return the MongoDB ID of the inserted document
    return result.inserted_id

def generate_book_data(data):
    writing_style = data.get('writing_style')
    book_description = data.get('book_description')
    chapter_titles = data.get('chapter_titles')
    chapter_elaborations = data.get('chapter_elaborations', [])

    if not all([writing_style, book_description, chapter_titles]):
        raise ValueError("Missing data for writing style, book description, or chapter titles")

    chapters_content = generate_book(writing_style, book_description, chapter_titles, chapter_elaborations)
    title = generate_title(book_description)

    file_stream = io.BytesIO()
    create_doc(title, 'Author Name', chapters_content, chapter_titles, file_stream)
    file_stream.seek(0)

    unique_filename = str(uuid.uuid4())
    document_id = save_book(file_stream, title, unique_filename)
    download_url = f"http://localhost:5001/download/{str(document_id)}"

    return title, download_url, document_id

def generate_book(writing_style, book_description, chapter_titles, chapter_elaborations):
    all_chapters_content = []

    for i, chapter_title in enumerate(chapter_titles):
        print(f"Generating content for chapter '{chapter_title}'...")
        subtitles_content = []

        for subtitle_index in range(1, 2):
            attempts, subtitle_generated, max_attempts = 0, False, 3

            while not subtitle_generated and attempts < max_attempts:
                subtitle_prompt = f"Subtitle {subtitle_index}: Provide a detailed analysis and insights for '{chapter_title}', as part of a book in the style of '{writing_style}'. Ensure this section contains at least 800 words."
                
                if i < len(chapter_elaborations) and chapter_elaborations[i]:
                    subtitle_prompt += f" Follow the instructions provided and include the following additional information after thinking it through intelligently. Add them in a way that is balanced and not excessive or abnormal: {chapter_elaborations[i]}"                
                subtitle_content = generate_text(subtitle_prompt, max_tokens=3000)

                if word_count(subtitle_content) >= 500:
                    subtitle_generated = True
                    subtitles_content.append(subtitle_content)
                else:
                    attempts += 1
                    print(f"Insufficient content for subtitle {subtitle_index} in chapter '{chapter_title}'. Attempting regeneration, attempt #{attempts}.")

        chapter_content = "\n\n".join(subtitles_content)
        all_chapters_content.append(chapter_content)
        time.sleep(2)  # Pause between chapters to avoid rate limiting

    print("Book content generation completed.")
    return all_chapters_content

def download_book(document_id):
    try:
        # Convert string to ObjectId
        object_id = ObjectId(document_id)
    except Exception as e:
        return None, f"Invalid MongoDB ObjectId: {str(e)}", 400

    try:
        # Fetch the document by its ID from MongoDB
        book_document = collection.find_one({"_id": object_id})
        if not book_document:
            return None, "File not found", 404
        
        # Create a BytesIO object from the binary data stored in MongoDB
        file_stream = io.BytesIO(book_document['content'])
        filename = f"{book_document['title']}.docx"
        return file_stream, filename, None
    except Exception as e:
        # Log the error to debug
        print(f"Error accessing MongoDB: {str(e)}")
        return None, str(e), 500

